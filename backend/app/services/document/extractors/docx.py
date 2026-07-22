from pathlib import Path

from docx import Document as DocxDocument

from app.services.document.extractors.base import BaseExtractor
from app.models.document import ExtractionResult


class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files using python-docx."""

    def extract(self, filepath: Path) -> ExtractionResult:
        if filepath.stat().st_size == 0:
            return ExtractionResult(content="", page_count=0)

        doc = DocxDocument(str(filepath))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # python-docx does not expose page count — estimate by paragraph count
        full_text = "\n".join(paragraphs)
        estimated_pages = max(1, len(paragraphs) // 20) if paragraphs else 0

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.extend(row_texts)

        full_text = "\n".join(paragraphs)
        return ExtractionResult(content=full_text, page_count=estimated_pages)
